from __future__ import annotations

import io
import re
import zipfile
from collections import defaultdict
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


SECTION_RE = re.compile(
    r"^([ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ]+)\.\s*([^ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ]+?)(?=(?:[ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ]+\.)|$)"
)
DOT_LEADER_RE = re.compile(r"\.{4,}")
KEY_VALUE_RE = re.compile(r"^\s*([가-힣A-Za-z0-9 _-]+)\s*[:：]\s*(.+?)\s*$")

SECTION_KEYWORDS: dict[str, tuple[str, ...]] = {
    "발명 동기": ("동기", "문제", "배경", "필요", "계기", "불편", "폭염"),
    "문제 해결": ("해결", "구성", "동작", "센서", "제어", "방식", "작동"),
    "선행 연구 고찰": ("선행", "연구", "조사", "이론", "기술", "비교", "고찰"),
    "발명품 제작": ("제작", "설계", "조립", "과정", "프로토타입", "실험", "구현"),
    "전망 및 활용성": ("활용", "효과", "기대", "전망", "확장", "적용", "장점"),
    "제작 후기": ("후기", "느낀", "배운", "아쉬운", "개선", "소감"),
    "참고 자료": ("참고", "자료", "출처", "링크", "문헌"),
}

CANONICAL_LABELS = {
    "작품명": "project_title",
    "제목": "project_title",
    "문서제목": "project_title",
    "대회명": "competition_title",
    "행사명": "competition_title",
    "출품학생": "student_name",
    "학생": "student_name",
    "지도교원": "advisor_name",
    "지도교사": "advisor_name",
    "출품분야": "category",
    "분야": "category",
    "날짜": "document_date",
    "작성일": "document_date",
}


@dataclass
class SectionTemplate:
    numeral: str
    title: str
    paragraph_count: int = 0
    table_count: int = 0
    image_slots: int = 0

    @property
    def full_title(self) -> str:
        return f"{self.numeral}. {self.title}"


@dataclass
class ReferenceAnalysis:
    source_name: str
    competition_title: str
    page_width_cm: float
    page_height_cm: float
    margin_top_cm: float
    margin_bottom_cm: float
    margin_left_cm: float
    margin_right_cm: float
    cover_fields: list[str] = field(default_factory=list)
    sections: list[SectionTemplate] = field(default_factory=list)
    total_tables: int = 0
    total_images: int = 0


@dataclass
class SummaryPayload:
    project_title: str = ""
    competition_title: str = ""
    student_name: str = ""
    advisor_name: str = ""
    category: str = ""
    document_date: str = ""
    raw_text: str = ""
    loose_paragraphs: list[str] = field(default_factory=list)
    section_text: dict[str, list[str]] = field(default_factory=lambda: defaultdict(list))
    references: list[str] = field(default_factory=list)


def _iter_docx_text(file_name: str, data: bytes) -> str:
    suffix = Path(file_name).suffix.lower()
    if suffix in {".txt", ".md"}:
        return data.decode("utf-8", errors="ignore")

    if suffix == ".docx":
        doc = Document(io.BytesIO(data))
        lines: list[str] = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                lines.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                line = " | ".join([cell for cell in cells if cell])
                if line:
                    lines.append(line)
        return "\n".join(lines)

    raise ValueError("지원하지 않는 요약서 형식입니다. TXT, MD, DOCX만 가능합니다.")


def analyze_reference_docx(data: bytes, source_name: str) -> ReferenceAnalysis:
    doc = Document(io.BytesIO(data))
    section = doc.sections[0]

    analysis = ReferenceAnalysis(
        source_name=source_name,
        competition_title=_extract_competition_title(doc),
        page_width_cm=_emu_to_cm(section.page_width),
        page_height_cm=_emu_to_cm(section.page_height),
        margin_top_cm=_emu_to_cm(section.top_margin),
        margin_bottom_cm=_emu_to_cm(section.bottom_margin),
        margin_left_cm=_emu_to_cm(section.left_margin),
        margin_right_cm=_emu_to_cm(section.right_margin),
        cover_fields=_extract_cover_fields(doc),
        total_tables=len(doc.tables),
        total_images=_count_total_images(doc),
    )

    current_section: SectionTemplate | None = None
    seen_titles: set[str] = set()
    body = doc.element.body

    for child in body.iterchildren():
        local_name = child.tag.rsplit("}", 1)[-1]

        if local_name == "p":
            text = "".join(child.itertext()).strip()
            heading = _match_heading(text)
            if heading and heading not in seen_titles:
                current_section = SectionTemplate(numeral=heading[0], title=heading[1])
                analysis.sections.append(current_section)
                seen_titles.add(heading)
                continue

            if current_section and text:
                current_section.paragraph_count += 1

            if current_section and _paragraph_has_drawing(child):
                current_section.image_slots += 1

        elif local_name == "tbl":
            if current_section:
                current_section.table_count += 1

    return analysis


def parse_summary_document(file_name: str, data: bytes, reference: ReferenceAnalysis) -> SummaryPayload:
    text = _iter_docx_text(file_name, data)
    payload = SummaryPayload(raw_text=text)

    current_section: str | None = None
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        heading = _normalize_summary_heading(line, reference.sections)
        if heading:
            current_section = heading
            continue

        match = KEY_VALUE_RE.match(line)
        if match:
            label = CANONICAL_LABELS.get(match.group(1).strip())
            if label:
                setattr(payload, label, match.group(2).strip())
                continue

        if current_section:
            payload.section_text[current_section].append(line)
        else:
            payload.loose_paragraphs.append(line)

    if not payload.project_title:
        payload.project_title = _guess_title_from_loose_paragraphs(payload.loose_paragraphs)
    if not payload.competition_title:
        payload.competition_title = reference.competition_title

    payload.references.extend(_extract_references(payload))
    return payload


def build_docx_document(reference: ReferenceAnalysis, summary: SummaryPayload) -> bytes:
    doc = Document()
    _apply_page_layout(doc, reference)
    _build_cover_page(doc, reference, summary)
    _build_toc_page(doc, reference)
    _build_body(doc, reference, summary)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def build_hwpx_document(reference: ReferenceAnalysis, summary: SummaryPayload, template_dir: Path) -> bytes:
    template_files = {
        relative_path: (template_dir / relative_path).read_bytes()
        for relative_path in [
            "mimetype",
            "version.xml",
            "settings.xml",
            "META-INF/container.xml",
            "META-INF/container.rdf",
            "META-INF/manifest.xml",
            "Contents/header.xml",
            "Preview/PrvImage.png",
        ]
    }

    content_hpf = _render_hwpx_content_hpf(summary.project_title or "문서")
    section_xml = _render_hwpx_section_xml(reference, summary)
    preview_text = _render_preview_text(reference, summary)

    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("mimetype", template_files["mimetype"], compress_type=zipfile.ZIP_STORED)
        for name, data in template_files.items():
            if name == "mimetype":
                continue
            zf.writestr(name, data)
        zf.writestr("Contents/content.hpf", content_hpf.encode("utf-8"))
        zf.writestr("Contents/section0.xml", section_xml.encode("utf-8"))
        zf.writestr("Preview/PrvText.txt", preview_text.encode("utf-8"))
    return output.getvalue()


def _extract_competition_title(doc: Document) -> str:
    for table in doc.tables[:2]:
        text = " ".join(cell.text.strip() for row in table.rows for cell in row.cells).strip()
        if "대회" in text or "경진대회" in text:
            cleaned = re.sub(r"\d{4}\.\s*\.\s*\.", "", text)
            cleaned = re.sub(r"\s+", " ", cleaned).strip()
            return cleaned
    return "문서 제작 기준안"


def _extract_cover_fields(doc: Document) -> list[str]:
    labels = []
    for table in doc.tables[:4]:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text in {"출품학생", "지도교원", "지도교사", "출품분야", "출품번호"}:
                    labels.append(text)
    unique_labels: list[str] = []
    for label in labels:
        if label not in unique_labels:
            unique_labels.append(label)
    return unique_labels


def _count_total_images(doc: Document) -> int:
    body = doc.element.body
    count = 0
    for child in body.iterchildren():
        if child.tag.rsplit("}", 1)[-1] == "p" and _paragraph_has_drawing(child):
            count += 1
    return count


def _paragraph_has_drawing(paragraph_element) -> bool:
    xml = paragraph_element.xml
    return "<w:drawing" in xml or "<w:pict" in xml


def _match_heading(text: str) -> tuple[str, str] | None:
    candidate = text.strip()
    if not candidate or DOT_LEADER_RE.search(candidate):
        return None

    match = SECTION_RE.match(candidate)
    if not match:
        return None

    numeral, title = match.group(1).strip(), match.group(2).strip()
    title = title.rstrip(".").strip()
    title = title.replace("  ", " ")
    return numeral, title


def _normalize_summary_heading(line: str, sections: Iterable[SectionTemplate]) -> str | None:
    stripped = line.strip().lstrip("#").strip()
    stripped = re.sub(r"^\d+\.\s*", "", stripped)

    heading_match = SECTION_RE.match(stripped)
    if heading_match:
        stripped = heading_match.group(2).strip()

    for section in sections:
        if stripped == section.title:
            return section.title
        if stripped.replace(" ", "") == section.title.replace(" ", ""):
            return section.title
    return None


def _guess_title_from_loose_paragraphs(paragraphs: list[str]) -> str:
    for paragraph in paragraphs:
        if len(paragraph) <= 40 and "http" not in paragraph.lower():
            return paragraph
    return paragraphs[0] if paragraphs else "새 작품 설명서"


def _extract_references(payload: SummaryPayload) -> list[str]:
    refs = []
    for paragraph in payload.loose_paragraphs:
        if "http://" in paragraph or "https://" in paragraph or "출처" in paragraph:
            refs.append(paragraph)
    if "참고 자료" in payload.section_text:
        refs.extend(payload.section_text["참고 자료"])
    return refs


def _apply_page_layout(doc: Document, reference: ReferenceAnalysis) -> None:
    section = doc.sections[0]
    section.page_width = Cm(reference.page_width_cm)
    section.page_height = Cm(reference.page_height_cm)
    section.top_margin = Cm(reference.margin_top_cm)
    section.bottom_margin = Cm(reference.margin_bottom_cm)
    section.left_margin = Cm(reference.margin_left_cm)
    section.right_margin = Cm(reference.margin_right_cm)


def _build_cover_page(doc: Document, reference: ReferenceAnalysis, summary: SummaryPayload) -> None:
    info_box = doc.add_table(rows=2, cols=1)
    info_box.style = "Table Grid"
    info_box.columns[0].width = Inches(1.2)
    info_box.rows[0].height = Cm(1.0)
    info_box.rows[1].height = Cm(1.5)

    label = info_box.cell(0, 0).paragraphs[0]
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = label.add_run("출품번호")
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(24)
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run(summary.competition_title or reference.competition_title)
    run.bold = True
    run.font.name = "HY헤드라인M"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "HY헤드라인M")
    run.font.size = Pt(22)

    project = doc.add_paragraph()
    project.alignment = WD_ALIGN_PARAGRAPH.CENTER
    project.paragraph_format.space_after = Pt(10)
    run = project.add_run(summary.project_title or "새 작품 설명서")
    run.bold = True
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(16)

    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_paragraph.add_run(summary.document_date or "2026.    .    .")
    date_run.font.name = "HY헤드라인M"
    date_run._element.rPr.rFonts.set(qn("w:eastAsia"), "HY헤드라인M")
    date_run.font.size = Pt(14)

    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = "Table Grid"
    info_table.rows[0].cells[0].text = "출품학생"
    info_table.rows[1].cells[0].text = "지도교원"
    info_table.rows[2].cells[0].text = "출품분야"
    info_table.rows[3].cells[0].text = "문서형식"
    info_table.rows[0].cells[1].text = summary.student_name or ""
    info_table.rows[1].cells[1].text = summary.advisor_name or ""
    info_table.rows[2].cells[1].text = summary.category or ""
    info_table.rows[3].cells[1].text = "기준 문서 결 반영"

    for row in info_table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Malgun Gothic"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
                    run.font.size = Pt(12)

    doc.add_page_break()


def _build_toc_page(doc: Document, reference: ReferenceAnalysis) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("차 례")
    run.bold = True
    run.font.name = "HYGungSo"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "한양신명조")
    run.font.size = Pt(20)

    for section in reference.sections:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(section.full_title)
        run.font.name = "Batang"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "한양신명조")
        run.font.size = Pt(11)

    doc.add_page_break()


def _build_body(doc: Document, reference: ReferenceAnalysis, summary: SummaryPayload) -> None:
    fallback_sections = _distribute_loose_paragraphs(summary, reference.sections)

    for index, section in enumerate(reference.sections):
        if index > 0:
            doc.add_section(WD_SECTION_START.NEW_PAGE)
            _apply_page_layout(doc, reference)

        heading = doc.add_paragraph()
        heading.paragraph_format.space_after = Pt(10)
        run = heading.add_run(section.full_title)
        run.bold = True
        run.font.name = "Batang"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "한양신명조")
        run.font.size = Pt(18)

        paragraphs = _compose_section_paragraphs(section.title, summary, fallback_sections.get(section.title, []))
        if not paragraphs:
            paragraphs = [f"{section.title}에 해당하는 내용을 요약서를 바탕으로 정리해 주세요."]

        for paragraph_text in paragraphs:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.line_spacing = 1.5
            run = paragraph.add_run(paragraph_text)
            run.font.name = "Gulim"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "굴림")
            run.font.size = Pt(11)

        for slot_index in range(max(section.image_slots, _default_photo_slots(section.title))):
            _add_photo_placeholder(doc, slot_index + 1)

    if summary.references and not any(section.title == "참고 자료" for section in reference.sections):
        doc.add_page_break()
        heading = doc.add_paragraph()
        run = heading.add_run("참고 자료")
        run.bold = True
        run.font.name = "Batang"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "한양신명조")
        run.font.size = Pt(18)
        for ref in summary.references:
            doc.add_paragraph(ref)


def _distribute_loose_paragraphs(summary: SummaryPayload, sections: Iterable[SectionTemplate]) -> dict[str, list[str]]:
    distributed: dict[str, list[str]] = {section.title: [] for section in sections}
    remaining = list(summary.loose_paragraphs)

    for section in sections:
        title = section.title
        if title in summary.section_text:
            continue

        keywords = SECTION_KEYWORDS.get(title, ())
        matched = [p for p in remaining if any(keyword in p for keyword in keywords)]
        if matched:
            distributed[title].extend(matched[:4])
            remaining = [p for p in remaining if p not in matched[:4]]

    for section in sections:
        if distributed[section.title]:
            continue
        if remaining:
            take = remaining[: min(3, len(remaining))]
            distributed[section.title].extend(take)
            remaining = remaining[len(take) :]

    return distributed


def _compose_section_paragraphs(title: str, summary: SummaryPayload, fallback: list[str]) -> list[str]:
    if title == "참고 자료":
        return summary.references or fallback

    paragraphs = list(summary.section_text.get(title, [])) or list(fallback)
    if paragraphs:
        return _normalize_paragraphs(paragraphs)

    intro = {
        "발명 동기": "이 작품은 일상에서 느낀 불편함과 해결 필요성에서 출발했습니다.",
        "문제 해결": "요약서에 제시된 핵심 아이디어를 바탕으로 해결 구조를 정리했습니다.",
        "선행 연구 고찰": "관련 개념과 기존 사례를 검토해 작품의 방향을 보완했습니다.",
        "발명품 제작": "구상, 설계, 제작, 점검 순서로 제작 과정을 정리했습니다.",
        "전망 및 활용성": "현장 적용 가능성과 기대 효과를 중심으로 정리했습니다.",
        "제작 후기": "제작 과정에서 느낀 점과 개선 방향을 함께 정리했습니다.",
    }
    return [intro.get(title, f"{title} 항목을 기준으로 내용을 정리했습니다.")]


def _normalize_paragraphs(paragraphs: list[str]) -> list[str]:
    normalized: list[str] = []
    buffer: list[str] = []

    for paragraph in paragraphs:
        text = paragraph.strip()
        if not text:
            continue
        if text.startswith(("-", "•", "*")):
            normalized.append(text.lstrip("-•* ").strip())
            continue
        buffer.append(text)

    if buffer:
        normalized.extend(_chunk_text(buffer))

    return normalized


def _chunk_text(lines: list[str], max_len: int = 210) -> list[str]:
    chunks: list[str] = []
    current = ""
    for line in lines:
        if len(current) + len(line) + 1 <= max_len:
            current = f"{current} {line}".strip()
        else:
            if current:
                chunks.append(current)
            current = line
    if current:
        chunks.append(current)
    return chunks


def _add_photo_placeholder(doc: Document, number: int) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell.width = Inches(5.8)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f"[사진 추가 예정 {number}]")
    run.bold = True
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(12)
    doc.add_paragraph()


def _default_photo_slots(title: str) -> int:
    if title == "발명품 제작":
        return 3
    if title in {"문제 해결", "전망 및 활용성"}:
        return 1
    return 0


def _render_hwpx_content_hpf(title: str) -> str:
    safe_title = escape(title)
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes" ?>'
        '<opf:package xmlns:ha="http://www.hancom.co.kr/hwpml/2011/app" '
        'xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph" '
        'xmlns:hp10="http://www.hancom.co.kr/hwpml/2016/paragraph" '
        'xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" '
        'xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" '
        'xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" '
        'xmlns:hhs="http://www.hancom.co.kr/hwpml/2011/history" '
        'xmlns:hm="http://www.hancom.co.kr/hwpml/2011/master-page" '
        'xmlns:hpf="http://www.hancom.co.kr/schema/2011/hpf" '
        'xmlns:dc="http://purl.org/dc/elements/1.1/" '
        'xmlns:opf="http://www.idpf.org/2007/opf/" '
        'xmlns:ooxmlchart="http://www.hancom.co.kr/hwpml/2016/ooxmlchart" '
        'xmlns:hwpunitchar="http://www.hancom.co.kr/hwpml/2016/HwpUnitChar" '
        'xmlns:epub="http://www.idpf.org/2007/ops" '
        'xmlns:config="urn:oasis:names:tc:opendocument:xmlns:config:1.0">'
        "<opf:metadata>"
        f"<opf:title>{safe_title}</opf:title>"
        "<opf:language>ko</opf:language>"
        '<opf:meta name="creator" content="text">Codex</opf:meta>'
        '<opf:meta name="lastsaveby" content="text">Codex</opf:meta>'
        "</opf:metadata>"
        "<opf:manifest>"
        '<opf:item id="header" href="Contents/header.xml" media-type="application/xml"/>'
        '<opf:item id="section0" href="Contents/section0.xml" media-type="application/xml"/>'
        '<opf:item id="settings" href="settings.xml" media-type="application/xml"/>'
        "</opf:manifest>"
        "<opf:spine>"
        '<opf:itemref idref="header" linear="yes"/>'
        '<opf:itemref idref="section0" linear="yes"/>'
        "</opf:spine>"
        "</opf:package>"
    )


def _render_preview_text(reference: ReferenceAnalysis, summary: SummaryPayload) -> str:
    lines = [summary.project_title or "새 작품 설명서", ""]
    for section in reference.sections:
        lines.append(section.full_title)
        for paragraph in _compose_section_paragraphs(section.title, summary, []):
            lines.append(paragraph)
        if max(section.image_slots, _default_photo_slots(section.title)) > 0:
            lines.append("[사진 추가 예정]")
        lines.append("")
    return "\n".join(lines).strip()


def _render_hwpx_section_xml(reference: ReferenceAnalysis, summary: SummaryPayload) -> str:
    paragraphs = [_hwpx_first_paragraph()]
    paragraphs.extend(_hwpx_plain_paragraph(summary.project_title or "새 작품 설명서", char_pr="9"))
    paragraphs.extend(_hwpx_plain_paragraph(summary.competition_title or reference.competition_title, char_pr="8"))
    paragraphs.extend(_hwpx_plain_paragraph(summary.document_date or "2026.    .    ."))
    paragraphs.extend(_hwpx_plain_paragraph(""))
    paragraphs.extend(_hwpx_plain_paragraph("차 례", char_pr="9"))

    for section in reference.sections:
        paragraphs.extend(_hwpx_plain_paragraph(section.full_title, char_pr="8"))

    paragraphs.extend(_hwpx_plain_paragraph(""))

    for section in reference.sections:
        paragraphs.extend(_hwpx_plain_paragraph(section.full_title, char_pr="9"))
        section_paragraphs = _compose_section_paragraphs(section.title, summary, [])
        for paragraph in section_paragraphs:
            paragraphs.extend(_hwpx_plain_paragraph(paragraph))
        for slot_index in range(max(section.image_slots, _default_photo_slots(section.title))):
            paragraphs.extend(_hwpx_plain_paragraph(f"[사진 추가 예정 {slot_index + 1}]"))
        paragraphs.extend(_hwpx_plain_paragraph(""))

    joined = "".join(paragraphs)
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes" ?>'
        '<hs:sec xmlns:ha="http://www.hancom.co.kr/hwpml/2011/app" '
        'xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph" '
        'xmlns:hp10="http://www.hancom.co.kr/hwpml/2016/paragraph" '
        'xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" '
        'xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" '
        'xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" '
        'xmlns:hhs="http://www.hancom.co.kr/hwpml/2011/history" '
        'xmlns:hm="http://www.hancom.co.kr/hwpml/2011/master-page" '
        'xmlns:hpf="http://www.hancom.co.kr/schema/2011/hpf" '
        'xmlns:dc="http://purl.org/dc/elements/1.1/" '
        'xmlns:opf="http://www.idpf.org/2007/opf/" '
        'xmlns:ooxmlchart="http://www.hancom.co.kr/hwpml/2016/ooxmlchart" '
        'xmlns:hwpunitchar="http://www.hancom.co.kr/hwpml/2016/HwpUnitChar" '
        'xmlns:epub="http://www.idpf.org/2007/ops" '
        'xmlns:config="urn:oasis:names:tc:opendocument:xmlns:config:1.0">'
        f"{joined}</hs:sec>"
    )


def _hwpx_first_paragraph() -> str:
    return (
        '<hp:p id="1" paraPrIDRef="20" styleIDRef="0" pageBreak="0" columnBreak="0" merged="0">'
        '<hp:run charPrIDRef="7">'
        '<hp:secPr id="" textDirection="HORIZONTAL" spaceColumns="1134" tabStop="8000" '
        'tabStopVal="4000" tabStopUnit="HWPUNIT" outlineShapeIDRef="1" memoShapeIDRef="0" '
        'textVerticalWidthHead="0" masterPageCnt="0">'
        '<hp:grid lineGrid="0" charGrid="0" wonggojiFormat="0"/>'
        '<hp:startNum pageStartsOn="BOTH" page="0" pic="0" tbl="0" equation="0"/>'
        '<hp:visibility hideFirstHeader="0" hideFirstFooter="0" hideFirstMasterPage="0" border="SHOW_ALL" '
        'fill="SHOW_ALL" hideFirstPageNum="0" hideFirstEmptyLine="0" showLineNumber="0"/>'
        '<hp:lineNumberShape restartType="0" countBy="0" distance="0" startNumber="0"/>'
        '<hp:pagePr landscape="WIDELY" width="59528" height="84186" gutterType="LEFT_ONLY">'
        '<hp:margin header="4252" footer="4252" gutter="0" left="8504" right="8504" top="5668" bottom="4252"/>'
        '</hp:pagePr>'
        '<hp:footNotePr><hp:autoNumFormat type="DIGIT" userChar="" prefixChar="" suffixChar=")" supscript="0"/>'
        '<hp:noteLine length="-1" type="SOLID" width="0.12 mm" color="#000000"/>'
        '<hp:noteSpacing betweenNotes="283" belowLine="567" aboveLine="850"/>'
        '<hp:numbering type="CONTINUOUS" newNum="1"/>'
        '<hp:placement place="EACH_COLUMN" beneathText="0"/></hp:footNotePr>'
        '<hp:endNotePr><hp:autoNumFormat type="DIGIT" userChar="" prefixChar="" suffixChar=")" supscript="0"/>'
        '<hp:noteLine length="14692344" type="SOLID" width="0.12 mm" color="#000000"/>'
        '<hp:noteSpacing betweenNotes="0" belowLine="567" aboveLine="850"/>'
        '<hp:numbering type="CONTINUOUS" newNum="1"/>'
        '<hp:placement place="END_OF_DOCUMENT" beneathText="0"/></hp:endNotePr>'
        '<hp:pageBorderFill type="BOTH" borderFillIDRef="1" textBorder="PAPER" headerInside="0" footerInside="0" fillArea="PAPER">'
        '<hp:offset left="1417" right="1417" top="1417" bottom="1417"/></hp:pageBorderFill>'
        '</hp:secPr>'
        "</hp:run>"
        '<hp:run charPrIDRef="7"><hp:t></hp:t></hp:run>'
        '<hp:linesegarray><hp:lineseg textpos="0" vertpos="0" vertsize="1000" textheight="1000" baseline="850" spacing="200" horzpos="0" horzsize="42520" flags="393216"/></hp:linesegarray>'
        "</hp:p>"
    )


def _hwpx_plain_paragraph(text: str, char_pr: str = "7") -> list[str]:
    escaped = escape(text)
    return [
        (
            f'<hp:p id="0" paraPrIDRef="21" styleIDRef="0" pageBreak="0" columnBreak="0" merged="0">'
            f'<hp:run charPrIDRef="{char_pr}"><hp:t>{escaped}</hp:t></hp:run>'
            '<hp:linesegarray><hp:lineseg textpos="0" vertpos="0" vertsize="1000" '
            'textheight="1000" baseline="850" spacing="600" horzpos="0" horzsize="42520" '
            'flags="393216"/></hp:linesegarray></hp:p>'
        )
    ]


def _emu_to_cm(value) -> float:
    if value is None:
        return 0.0
    return round(value.cm, 2)
